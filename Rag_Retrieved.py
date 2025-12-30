from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain_community.document_loaders import Docx2txtLoader
from pathlib import Path
from langchain_openai import ChatOpenAI,OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
import os
import pandas as pd
from dotenv import load_dotenv
import re
from typing import List, Dict, Any, Tuple

import faiss
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch

# 尝试导入python-docx，用于更精确的文档结构处理
try:
    from docx import Document as DocxDocument
except ImportError:
    print("Warning: python-docx not found, falling back to Docx2txtLoader")
    DocxDocument = None

from helper_functions import (
    num_tokens_from_string, replace_t_with_space, replace_double_lines_with_one_line, 
    split_into_chapters, analyse_metric_results, escape_quotes, text_wrap, 
    extract_book_quotes_as_documents
)
from bm25_retriever import BM25Retriever
from reranker import Reranker
from reasoning_trace import trace_manager

load_dotenv()
os.environ["DEEPSEEK_API_KEY"] = os.getenv('DEEPSEEK_API_KEY')

def encode_excel(path, chunk_size=1000, chunk_overlap=200):
    """
    利用BGE-large-zh嵌入将Excel文件编码到矢量存储中。

    Args：
        path：指向Excel文件的路径。
        chunk_size：每个文本块的期望大小。
        chunk_overlap：连续区块之间的重叠程度。

    Returns：
        一个包含编码Excel内容的FAISS向量存储器。
    """

    # Load Excel documents
    # 使用pandas读取Excel文件
    df = pd.read_excel(path)

    # 将DataFrame转换为文档
    documents = []
    for index, row in df.iterrows():
        # 构建完整的文本内容
        row_text = f"案由名称：{row['案由名称']}\n"
        row_text += f"违法依据：{row['违法依据']}\n"
        row_text += f"行政处罚依据：{row['行政处罚依据']}\n"
        row_text += f"处罚内容：{row['处罚内容']}"
        
        # 构建丰富的元数据
        metadata = {
            "row_index": index,
            "source": path,
            "案由名称": row['案由名称'],
            "所属专业": row['所属专业'],
            "适用的法律法规": row['适用的法律法规'],
            "类型": "案由"
        }

        doc = Document(
            page_content=row_text.strip(),
            metadata=metadata
        )
        documents.append(doc)

    document_cleaned = replace_t_with_space(documents)

    # Create embeddings and vector store
    # 使用本地BGE-large-zh模型
    embeddings = HuggingFaceEmbeddings(
        model_name="local_models/bge-large-zh",
        model_kwargs={'device': 'cpu'}
    )   
    vectorstore = FAISS.from_documents(document_cleaned, embeddings)

    return vectorstore


def load_docx_with_structure(path: str) -> List[Dict[str, Any]]:
    """
    加载Word文档，保留文档结构信息
    
    Args:
        path: Word文档路径
        
    Returns:
        包含文档结构信息的段落列表，每个段落包含文本和元数据
    """
    # 检查文件扩展名
    file_ext = Path(path).suffix.lower()
    
    if file_ext == ".doc":
        # .doc文件使用Docx2txtLoader
        loader = Docx2txtLoader(path)
        documents = loader.load()
        text = documents[0].page_content
    elif file_ext == ".docx":
        if not DocxDocument:
            # 如果没有python-docx，使用Docx2txtLoader作为备选
            loader = Docx2txtLoader(path)
            documents = loader.load()
            text = documents[0].page_content
        else:
            # .docx文件使用python-docx
            doc = DocxDocument(path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
    else:
        # 其他格式使用Docx2txtLoader
        loader = Docx2txtLoader(path)
        documents = loader.load()
        text = documents[0].page_content
    
    # 无论使用哪种加载方式，都使用正则表达式分析文档结构
    structured_paragraphs = []
    
    # 识别章节标题的正则表达式
    chapter_pattern = re.compile(r'^(第[一二三四五六七八九十百千]+[章节条款项])|^(第一章|第二章|第三章|第四章|第五章|第六章|第七章|第八章|第九章|第十章)|^([0-9]+\.)')
    
    # 按行分割文本
    lines = text.split('\n')
    
    current_chapter = ""
    current_section = ""
    current_paragraph = ""
    
    for line in lines:
        line = line.strip()
        if not line:
            # 如果当前有未处理的段落，先保存
            if current_paragraph:
                structured_paragraphs.append({
                    "text": current_paragraph,
                    "metadata": {
                        "type": "paragraph",
                        "level": 0,
                        "chapter": current_chapter,
                        "section": current_section
                    }
                })
                current_paragraph = ""
            continue
        
        # 检查是否为章节标题
        if chapter_pattern.match(line):
            # 如果当前有未处理的段落，先保存
            if current_paragraph:
                structured_paragraphs.append({
                    "text": current_paragraph,
                    "metadata": {
                        "type": "paragraph",
                        "level": 0,
                        "chapter": current_chapter,
                        "section": current_section
                    }
                })
                current_paragraph = ""
            
            # 更新当前章节信息
            if line.startswith("第") and ("章" in line or "节" in line):
                current_chapter = line
                current_section = ""
                structured_paragraphs.append({
                    "text": line,
                    "metadata": {
                        "type": "chapter",
                        "level": 1,
                        "chapter": line
                    }
                })
            else:
                current_section = line
                structured_paragraphs.append({
                    "text": line,
                    "metadata": {
                        "type": "section",
                        "level": 2,
                        "chapter": current_chapter,
                        "section": line
                    }
                })
        else:
            # 普通段落，累积到当前段落
            current_paragraph += line + "\n"
    
    # 保存最后一个段落
    if current_paragraph:
        structured_paragraphs.append({
            "text": current_paragraph,
            "metadata": {
                "type": "paragraph",
                "level": 0,
                "chapter": current_chapter,
                "section": current_section
            }
        })
    
    return structured_paragraphs

def adaptive_chunk_split(structured_paragraphs: List[Dict[str, Any]], 
                        min_chunk_size: int = 500, 
                        max_chunk_size: int = 2000, 
                        overlap: int = 100) -> List[Document]:
    """
    自适应Chunk划分算法，根据文档结构动态调整Chunk大小
    
    Args:
        structured_paragraphs: 包含结构信息的段落列表
        min_chunk_size: 最小Chunk大小
        max_chunk_size: 最大Chunk大小
        overlap: 相邻Chunk的重叠字符数
        
    Returns:
        划分后的Document列表
    """
    chunks = []
    current_chunk = ""
    current_metadata = {}
    chunk_index = 0
    
    for para in structured_paragraphs:
        para_text = para["text"]
        para_metadata = para["metadata"]
        
        # 章节标题单独作为Chunk
        if para_metadata["type"] in ["chapter", "section"]:
            # 先保存当前Chunk
            if current_chunk:
                chunks.append(Document(
                    page_content=current_chunk.strip(),
                    metadata={**current_metadata, "chunk_index": chunk_index, "chunk_type": "content"}
                ))
                chunk_index += 1
                current_chunk = ""
            
            # 章节标题作为单独Chunk
            chunks.append(Document(
                page_content=para_text,
                metadata={**para_metadata, "chunk_index": chunk_index, "chunk_type": "heading"}
            ))
            chunk_index += 1
            current_metadata = para_metadata
        else:
            # 普通段落，根据大小决定是否合并
            if not current_chunk:
                # 新的Chunk开始
                current_chunk = para_text
                current_metadata = para_metadata
            else:
                # 检查合并后的大小
                combined_size = len(current_chunk) + len(para_text)
                if combined_size <= max_chunk_size:
                    # 合并到当前Chunk
                    current_chunk += "\n" + para_text
                else:
                    # 当前Chunk达到最大大小，保存并开始新Chunk
                    # 添加重叠部分
                    if len(current_chunk) > overlap:
                        overlap_text = current_chunk[-overlap:]
                    else:
                        overlap_text = current_chunk
                    
                    chunks.append(Document(
                        page_content=current_chunk.strip(),
                        metadata={**current_metadata, "chunk_index": chunk_index, "chunk_type": "content"}
                    ))
                    chunk_index += 1
                    
                    # 开始新Chunk，包含重叠部分
                    current_chunk = overlap_text + "\n" + para_text
                    current_metadata = para_metadata
    
    # 保存最后一个Chunk
    if current_chunk:
        chunks.append(Document(
            page_content=current_chunk.strip(),
            metadata={**current_metadata, "chunk_index": chunk_index, "chunk_type": "content"}
        ))
    
    return chunks

def generate_summary(text: str, max_length: int = 200) -> str:
    """
    生成文本摘要
    
    Args:
        text: 原始文本
        max_length: 摘要最大长度
        
    Returns:
        生成的摘要
    """
    # 简单实现：取文本的前max_length个字符，后续可以替换为更复杂的摘要模型
    if len(text) <= max_length:
        return text
    return text[:max_length] + "..."

def create_hierarchical_index(documents: List[Document], embeddings) -> Tuple[FAISS, FAISS]:
    """
    创建分层索引，分为摘要层和细节层
    
    Args:
        documents: 原始文档列表
        embeddings: 嵌入模型
        
    Returns:
        摘要层索引和细节层索引
    """
    # 生成摘要层
    summary_docs = []
    detail_docs = []
    
    for doc in documents:
        # 添加到细节层
        detail_docs.append(doc)
        
        # 生成摘要，添加到摘要层
        if doc.metadata.get("chunk_type") == "content":
            summary = generate_summary(doc.page_content)
            summary_docs.append(Document(
                page_content=summary,
                metadata={**doc.metadata, "index_type": "summary"}
            ))
        else:
            # 标题直接添加到摘要层
            summary_docs.append(Document(
                page_content=doc.page_content,
                metadata={**doc.metadata, "index_type": "summary"}
            ))
    
    # 创建索引
    summary_index = FAISS.from_documents(summary_docs, embeddings)
    detail_index = FAISS.from_documents(detail_docs, embeddings)
    
    return summary_index, detail_index

def encode_word(path, chunk_size=1000, chunk_overlap=200, hierarchical_index=False):
    """
    利用BGE-large-zh嵌入将Word文件编码到矢量存储中，支持自适应Chunk划分和分层索引。

    Args：
        path：指向Word文件的路径。
        chunk_size：每个文本块的期望大小。
        chunk_overlap：连续区块之间的重叠程度。
        hierarchical_index：是否创建分层索引

    Returns：
        一个包含编码Word内容的FAISS向量存储器，或（摘要层索引，细节层索引）的元组。
    """

    print(f"读取文档: {path}")
    
    # 加载带结构的文档
    structured_paragraphs = load_docx_with_structure(path)
    print(f"成功读取Word文档，识别到 {len(structured_paragraphs)} 个结构化段落")
    
    # 自适应Chunk划分
    chunks = adaptive_chunk_split(
        structured_paragraphs,
        min_chunk_size=chunk_size//2,
        max_chunk_size=chunk_size,
        overlap=chunk_overlap
    )
    print(f"文档自适应划分为 {len(chunks)} 个块")
    
    # 为每个块添加丰富的元数据
    for i, chunk in enumerate(chunks):
        # 添加文档级别元数据
        chunk.metadata.update({
            "类型": "法律法规",
            "来源": Path(path).stem,
            "文档路径": str(path),
            "总块数": len(chunks)
        })
    
    # 重新创建向量存储
    # 使用本地BGE-large-zh模型
    embeddings = HuggingFaceEmbeddings(
        model_name="local_models/bge-large-zh",
        model_kwargs={'device': 'cpu'}
    )   
    
    if hierarchical_index:
        # 创建分层索引
        summary_index, detail_index = create_hierarchical_index(chunks, embeddings)
        print(f"创建分层索引成功：摘要层 {len(summary_index.docstore._dict)} 个文档，细节层 {len(detail_index.docstore._dict)} 个文档")
        return summary_index, detail_index
    else:
        # 创建单层索引
        vectorstore = FAISS.from_documents(chunks, embeddings)
        print(f"创建单层索引成功：{len(vectorstore.docstore._dict)} 个文档")
        return vectorstore

def Creating_Vector_Store():
    """
    创建并保存向量存储。
    """
    file_path = ('data/案由清单.xlsx') # insert the path of the excel file
    word_path = ('data/最新版《行政处罚法》.docx')
    vectorstore = encode_excel(file_path)
    word_vectorstore = encode_word(word_path)
    vectorstore.save_local("case_causes_vector_store")
    word_vectorstore.save_local("legal_vector_store")

def create_retrievers():
    # 使用本地BGE-large-zh模型
    embeddings = HuggingFaceEmbeddings(
        model_name="local_models/bge-large-zh",
        model_kwargs={'device': 'cpu'}
    )   
    case_causes_vector_store = FAISS.load_local("case_causes_vector_store", embeddings, allow_dangerous_deserialization=True)
    legal_vector_store = FAISS.load_local("legal_vector_store", embeddings, allow_dangerous_deserialization=True)

    # 设置更大的k值以获取更多候选结果，用于后续混合检索
    case_causes_retriever = case_causes_vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 20})
    legal_retriever = legal_vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 20})

    return case_causes_retriever, legal_retriever

def _match_metadata(metadata, metadata_filters):
    """
    检查文档元数据是否匹配过滤条件
    
    Args：
        metadata：文档的元数据
        metadata_filters：元数据过滤条件，格式为{"key": "value"}
        
    Returns：
        如果匹配返回True，否则返回False
    """
    if not metadata_filters:
        return True
    
    for key, value in metadata_filters.items():
        if key not in metadata or metadata[key] != value:
            return False
    
    return True

def hybrid_retrieval(query, semantic_retriever, bm25_retriever, metadata_filters=None, k=10, session_id=None):
    """
    混合检索：结合语义检索和BM25关键词检索
    
    Args：
        query：查询文本
        semantic_retriever：语义检索器
        bm25_retriever：BM25关键词检索器
        metadata_filters：元数据过滤条件
        k：返回结果数量
        session_id：会话ID，用于记录推理轨迹
        
    Returns：
        混合检索结果
    """
    # 记录语义检索开始
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="开始语义检索",
            data={"query": query}
        )
    
    # 语义检索
    # 使用_langchain_core.Retriever的invoke方法，这是标准接口
    semantic_results = semantic_retriever.invoke(query)
    
    # 记录语义检索结果
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="语义检索完成",
            data={
                "semantic_results_count": len(semantic_results),
                "top_semantic_results": [
                    {"content": result.page_content[:100], "metadata": result.metadata}
                    for result in semantic_results[:3]
                ]
            }
        )
    
    # 记录BM25检索开始
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="开始BM25关键词检索",
            data={"query": query}
        )
    
    # BM25检索
    bm25_results_with_scores = bm25_retriever.retrieve(query, k=k, session_id=session_id)
    bm25_results = [doc for doc, score in bm25_results_with_scores]
    
    # 记录BM25检索结果
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="BM25检索完成",
            data={
                "bm25_results_count": len(bm25_results),
                "top_bm25_results": [
                    {"content": result.page_content[:100], "metadata": result.metadata, "score": score}
                    for result, score in bm25_results_with_scores[:3]
                ]
            }
        )
    
    # 元数据过滤
    if metadata_filters:
        if session_id:
            trace_manager.add_step(
                session_id=session_id,
                step_type="retrieval",
                description="应用元数据过滤",
                data={"metadata_filters": metadata_filters}
            )
        semantic_results = [doc for doc in semantic_results if _match_metadata(doc.metadata, metadata_filters)]
        bm25_results = [doc for doc in bm25_results if _match_metadata(doc.metadata, metadata_filters)]
        
        if session_id:
            trace_manager.add_step(
                session_id=session_id,
                step_type="retrieval",
                description="元数据过滤完成",
                data={
                    "filtered_semantic_results_count": len(semantic_results),
                    "filtered_bm25_results_count": len(bm25_results)
                }
            )
    
    # 合并结果（去重）
    combined_results = []
    seen_ids = set()
    
    # 先添加语义检索结果
    for doc in semantic_results:
        doc_id = f"{doc.metadata.get('source', '')}_{doc.metadata.get('row_index', '')}_{doc.metadata.get('chunk_index', '')}"
        if doc_id not in seen_ids:
            seen_ids.add(doc_id)
            combined_results.append((doc, 'semantic'))
    
    # 再添加BM25检索结果
    for doc in bm25_results:
        doc_id = f"{doc.metadata.get('source', '')}_{doc.metadata.get('row_index', '')}_{doc.metadata.get('chunk_index', '')}"
        if doc_id not in seen_ids:
            seen_ids.add(doc_id)
            combined_results.append((doc, 'bm25'))
    
    # 记录结果合并完成
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="结果合并完成",
            data={
                "combined_results_count": len(combined_results),
                "unique_documents_count": len(seen_ids)
            }
        )
    
    return combined_results

def retrieve_context_per_question(state):
    """
    根据问题从向量存储中检索上下文。

    Args：
        state：包含问题的状态字典。

    Returns：
        更新后的状态字典，包含检索到的上下文。
    """
    question = state["question"]
    retriever = state["retriever"]
    if "retrieved_context" not in state:
        state["retrieved_context"] = []
    if retriever == case_causes_retriever:
        case_causes_context = retriever.invoke(question)
        state["retrieved_context"] = case_causes_context
    elif retriever == legal_retriever:
        legal_context = retriever.invoke(question)
        state["retrieved_context"] = legal_context
    return state

def create_hybrid_retriever():
    """
    创建混合检索器，包括语义检索器、BM25检索器和重排序器
    
    Returns：
        语义检索器、BM25检索器和重排序器
    """
    # 加载向量存储
    embeddings = HuggingFaceEmbeddings(
        model_name="local_models/bge-large-zh",
        model_kwargs={'device': 'cpu'}
    )   
    case_causes_vector_store = FAISS.load_local("case_causes_vector_store", embeddings, allow_dangerous_deserialization=True)
    
    # 创建语义检索器
    semantic_retriever = case_causes_vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 20})
    
    # 重新加载文档以初始化BM25
    df = pd.read_excel('data/案由清单.xlsx')
    documents = []
    for index, row in df.iterrows():
        # 构建文档
        row_text = f"案由名称：{row['案由名称']}\n"
        row_text += f"违法依据：{row['违法依据']}\n"
        row_text += f"行政处罚依据：{row['行政处罚依据']}\n"
        row_text += f"处罚内容：{row['处罚内容']}"
        
        metadata = {
            "row_index": index,
            "source": 'data/案由清单.xlsx',
            "案由名称": row['案由名称'],
            "所属专业": row['所属专业'],
            "适用的法律法规": row['适用的法律法规'],
            "类型": "案由"
        }
        
        doc = Document(page_content=row_text.strip(), metadata=metadata)
        documents.append(doc)
    
    # 创建BM25检索器
    bm25_retriever = BM25Retriever(documents)
    
    # 创建重排序器
    reranker = Reranker()
    
    return semantic_retriever, bm25_retriever, reranker

def hierarchical_search(query, summary_index, detail_index, reranker, metadata_filters=None, k=5, session_id=None):
    """
    分层检索流程：先检索摘要层，再根据结果深入细节层
    
    Args：
        query：查询文本
        summary_index：摘要层索引
        detail_index：细节层索引
        reranker：重排序器
        metadata_filters：元数据过滤条件
        k：返回结果数量
        session_id：会话ID，用于记录推理轨迹
        
    Returns：
        重排序后的检索结果
    """
    # 记录分层检索开始
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="开始分层检索",
            data={"query": query, "metadata_filters": metadata_filters}
        )
    
    # 1. 先检索摘要层，获取相关章节
    summary_retriever = summary_index.as_retriever(search_type="similarity", search_kwargs={"k": 10})
    summary_results = summary_retriever.invoke(query)
    
    # 记录摘要层检索结果
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="摘要层检索完成",
            data={
                "summary_results_count": len(summary_results),
                "top_summary_results": [
                    {"content": result.page_content[:100], "metadata": result.metadata}
                    for result in summary_results[:3]
                ]
            }
        )
    
    # 2. 提取相关章节信息
    relevant_chapters = set()
    for result in summary_results:
        if "chapter" in result.metadata:
            relevant_chapters.add(result.metadata["chapter"])
    
    # 记录相关章节提取结果
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="提取相关章节信息",
            data={"relevant_chapters": list(relevant_chapters)}
        )
    
    # 3. 构建细节层检索的元数据过滤条件
    detail_filters = metadata_filters.copy() if metadata_filters else {}
    if relevant_chapters:
        # 只检索相关章节的内容
        detail_filters["chapter"] = list(relevant_chapters)
    
    # 记录细节层检索条件
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="构建细节层检索条件",
            data={"detail_filters": detail_filters}
        )
    
    # 4. 检索细节层
    detail_retriever = detail_index.as_retriever(search_type="similarity", search_kwargs={"k": 20})
    detail_results = detail_retriever.invoke(query)
    
    # 记录细节层检索结果
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="细节层检索完成",
            data={"detail_results_count": len(detail_results)}
        )
    
    # 5. 元数据过滤
    filtered_results = []
    for doc in detail_results:
        match = True
        for key, value in detail_filters.items():
            if key not in doc.metadata:
                match = False
                break
            if isinstance(value, list):
                if doc.metadata[key] not in value:
                    match = False
                    break
            else:
                if doc.metadata[key] != value:
                    match = False
                    break
        if match:
            filtered_results.append(doc)
    
    # 记录元数据过滤结果
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="细节层元数据过滤完成",
            data={
                "filtered_results_count": len(filtered_results),
                "original_results_count": len(detail_results)
            }
        )
    
    # 6. 重排序
    reranked_results = reranker.rerank(query, filtered_results, k=k, session_id=session_id)
    
    # 记录分层检索完成
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="分层检索完成",
            data={
                "final_results_count": len(reranked_results),
                "retrieved_chapters": list(relevant_chapters)
            }
        )
    
    return reranked_results

def hybrid_search(query, semantic_retriever, bm25_retriever, reranker, metadata_filters=None, k=5, hierarchical_index=None, session_id=None):
    """
    完整的混合检索流程：语义检索 + BM25检索 + 结果去重 + 重排序
    支持分层索引检索
    
    Args：
        query：查询文本
        semantic_retriever：语义检索器
        bm25_retriever：BM25检索器
        reranker：重排序器
        metadata_filters：元数据过滤条件
        k：返回结果数量
        hierarchical_index：分层索引元组 (summary_index, detail_index)
        session_id：会话ID，用于记录推理轨迹
        
    Returns：
        重排序后的检索结果
    """
    # 记录检索开始
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="开始混合检索",
            data={
                "query": query,
                "metadata_filters": metadata_filters,
                "k": k,
                "use_hierarchical_index": bool(hierarchical_index)
            }
        )
    
    # 如果提供了分层索引，使用分层检索
    if hierarchical_index:
        summary_index, detail_index = hierarchical_index
        if session_id:
            trace_manager.add_step(
                session_id=session_id,
                step_type="retrieval",
                description="使用分层检索策略",
                data={}
            )
        return hierarchical_search(query, summary_index, detail_index, reranker, metadata_filters, k, session_id)
    
    # 否则使用传统混合检索
    combined_results = hybrid_retrieval(query, semantic_retriever, bm25_retriever, metadata_filters, k=20, session_id=session_id)
    
    # 提取文档列表
    documents = [doc for doc, source in combined_results]
    
    # 重排序
    reranked_results = reranker.rerank(query, documents, k=k, session_id=session_id)
    
    # 记录检索完成
    if session_id:
        trace_manager.add_step(
            session_id=session_id,
            step_type="retrieval",
            description="混合检索完成",
            data={
                "total_documents": len(documents),
                "reranked_results_count": len(reranked_results),
                "retrieved_sources": list(set([doc.metadata.get('来源', 'unknown') for doc in reranked_results]))
            }
        )
    
    return reranked_results


if __name__=="__main__":
    # 首先创建向量存储
    # Creating_Vector_Store()
    
    # 测试新的法律文档处理和索引优化功能
    print("=== 测试法律文档处理和索引优化 ===")
    
    # 测试文件路径 - 只处理.docx文件
    test_files = [
        "data/医疗机构管理条例_20220329.docx",
        "data/最新版《行政处罚法》.docx"
    ]
    
    # 测试自适应Chunk划分和分层索引
    for file_path in test_files:
        if os.path.exists(file_path):
            print(f"\n--- 处理文件：{file_path} ---")
            
            # 创建分层索引
            print("创建分层索引...")
            try:
                result = encode_word(file_path, hierarchical_index=True)
                if isinstance(result, tuple):
                    summary_index, detail_index = result
                else:
                    print(f"警告：{file_path} 未返回分层索引")
                    continue
            
                # 测试分层检索
                print("测试分层检索...")
                from reranker import Reranker
                reranker = Reranker()
                
                # 测试查询
                if "行政处罚法" in file_path:
                    test_query = "行政处罚的种类有哪些？"
                else:
                    test_query = "医疗机构执业许可证的申请条件是什么？"
                
                results = hierarchical_search(
                    query=test_query,
                    summary_index=summary_index,
                    detail_index=detail_index,
                    reranker=reranker,
                    k=3
                )
                
                # 输出结果
                print(f"查询：{test_query}")
                print(f"检索到 {len(results)} 个相关结果：")
                for i, result in enumerate(results):
                    print(f"\n结果 {i+1}：")
                    print(f"内容：{result.page_content[:200]}...")
                    print(f"元数据：{result.metadata}")
                
                print("-" * 50)
            except Exception as e:
                print(f"处理文件 {file_path} 时出错：{str(e)}")
                print("-" * 50)
    
    # 测试传统混合检索器
    print("\n=== 测试传统混合检索器 ===")
    semantic_retriever, bm25_retriever, reranker = create_hybrid_retriever()
    
    # 执行混合检索
    query = "未取得医疗机构执业许可证擅自执业"
    results = hybrid_search(query, semantic_retriever, bm25_retriever, reranker, k=3)
    
    # 带元数据过滤的混合检索
    metadata_filters = {"所属专业": "医疗机构"}
    results = hybrid_search(query, semantic_retriever, bm25_retriever, reranker, metadata_filters=metadata_filters, k=1)
    print(results)


