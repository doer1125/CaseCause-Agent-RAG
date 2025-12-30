import torch
import time
from transformers import AutoModelForCausalLM, AutoTokenizer
from peft import AutoPeftModelForCausalLM
import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

# 设置CUDA设备
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(f"Using device: {device}")

class HealthDocumentGenerator:
    def __init__(self, model_path):
        self.model_path = model_path
        self.model = None
        self.tokenizer = None
        self.load_model()
    
    def load_model(self):
        """加载微调后的模型"""
        print(f"Loading fine-tuned model from {self.model_path}...")
        
        # 加载微调后的模型 - 优化GPU利用率
        self.model = AutoPeftModelForCausalLM.from_pretrained(
            self.model_path,
            device_map="balanced",  # 更平衡的设备映射
            trust_remote_code=True,
            dtype=torch.bfloat16,  # 使用新的dtype参数替代torch_dtype
            offload_folder=None,  # 关闭CPU offload，提高速度
            offload_buffers=False
        )
        
        # 尝试将模型完全移到GPU上
        self.model = self.model.to(device)
        
        self.tokenizer = AutoTokenizer.from_pretrained(
            self.model_path,
            trust_remote_code=True,
            padding_side="right"
        )
        
        # 优化模型推理性能
        self.model = self.model.eval()
        
        print("Model loaded successfully!")
    
    def generate_document(self, instruction):
        """根据用户指令生成公文"""
        print(f"Generating document for instruction: {instruction}")
        
        # 构建聊天模板
        messages = [
            {"role": "system", "content": "你是一名专业的卫健委公文撰写专家，请严格按照公文规范格式和术语要求撰写文书。"},
            {"role": "user", "content": instruction}
        ]
        
        # 应用聊天模板
        text = self.tokenizer.apply_chat_template(
            messages,
            tokenize=False,
            add_generation_prompt=True
        )
        
        # 编码文本
        inputs = self.tokenizer(
            text,
            return_tensors="pt",
            truncation=True,
            max_length=2048
        )
        
        # 将input_ids移动到与模型相同的设备上
        model_device = next(self.model.parameters()).device
        inputs = {k: v.to(model_device) for k, v in inputs.items()}
        
        # 生成文本 - 优化生成速度
        with torch.no_grad():
            outputs = self.model.generate(
                **inputs,
                max_new_tokens=512,  # 减少生成的token数量
                temperature=0.3,      # 降低随机性，提高生成速度
                top_p=0.8,            # 减少候选词数量
                repetition_penalty=1.0,
                do_sample=False,      # 使用确定性生成，速度更快
                num_beams=1,          # 不使用beam search
                use_cache=True        # 启用缓存加速生成
            )
        
        # 解码并输出结果 - 保留特殊标记以便精确提取
        generated_text = self.tokenizer.decode(
            outputs[0],
            skip_special_tokens=False,
            clean_up_tokenization_spaces=True
        )
        
        # 提取助手的回复 - 更严格的提取逻辑
        assistant_response = ""
        
        # 先找到assistant开始标记
        assistant_start = generated_text.find("<|im_start|>assistant")
        if assistant_start != -1:
            # 找到assistant开始标记，提取之后的内容
            assistant_content = generated_text[assistant_start + len("<|im_start|>assistant"):]
            
            # 找到对应的end标记
            assistant_end = assistant_content.find("<|im_end|>")
            if assistant_end != -1:
                # 有end标记，提取中间内容
                assistant_response = assistant_content[:assistant_end].strip()
            else:
                # 没有end标记，提取全部
                assistant_response = assistant_content.strip()
        else:
            # 没有找到assistant标记，使用skip_special_tokens=True重新解码
            assistant_response = self.tokenizer.decode(
                outputs[0],
                skip_special_tokens=True,
                clean_up_tokenization_spaces=True
            )
        
        # 进一步清理：移除所有可能的提示词和标记
        # 移除system提示
        if "system\n" in assistant_response:
            assistant_response = assistant_response.split("system\n", 1)[-1].strip()
        # 移除user提示
        if "user\n" in assistant_response:
            assistant_response = assistant_response.split("user\n", 1)[-1].strip()
        # 移除assistant标记
        if "assistant\n" in assistant_response:
            assistant_response = assistant_response.split("assistant\n", 1)[-1].strip()
        # 移除行首的角色标记
        lines = assistant_response.split("\n")
        clean_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                clean_lines.append(line)
            elif not (line.startswith("system") or line.startswith("user") or line.startswith("assistant")):
                clean_lines.append(line)
        assistant_response = "\n".join(clean_lines).strip()
        
        return assistant_response
    
    def save_to_pdf(self, content, output_path):
        """将生成的公文保存为PDF：先创建Word文档，再转换为PDF"""
        print(f"Saving document to PDF: {output_path}")
        
        # 使用唯一的临时文件名，避免冲突
        temp_dir = os.path.dirname(output_path) if os.path.dirname(output_path) else "."
        temp_filename = f"temp_{os.getpid()}_{int(time.time() * 1000)}.docx"
        word_temp = os.path.join(temp_dir, temp_filename)
        
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # 添加标题
            title_paragraph = doc.add_paragraph("卫健委公文")
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.runs[0]
            title_run.font.name = "微软雅黑"
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            
            # 添加空行
            doc.add_paragraph()
            
            # 处理生成的内容
            lines = content.split('\n')
            title_found = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                # 检查是否为标题（通常以"关于"开头）
                if line.startswith("关于") and ":" not in line and not title_found:
                    # 是标题，特殊处理
                    title_para = doc.add_paragraph(line)
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title_para.runs[0]
                    title_run.font.name = "微软雅黑"
                    title_run.font.size = Pt(18)
                    title_run.font.bold = True
                    doc.add_paragraph()
                    title_found = True
                else:
                    # 是正文内容
                    para = doc.add_paragraph(line)
                    for run in para.runs:
                        run.font.name = "仿宋"
                        run.font.size = Pt(14)
                    # 设置段落对齐和行距
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(8)
            
            # 保存Word文档
            doc.save(word_temp)
            print(f"Word document saved temporarily: {word_temp}")
            
            # 确保文件已正确保存
            if not os.path.exists(word_temp):
                raise FileNotFoundError(f"Temporary Word file not found: {word_temp}")
            
            # 将Word转换为PDF
            convert(word_temp, output_path)
            print(f"PDF saved successfully: {output_path}")
            
        except Exception as e:
            print(f"❌ PDF生成失败：{str(e)}")
            # 尝试使用不同的方式生成
            print("尝试直接生成PDF...")
            self._direct_pdf_generation(content, output_path)
            return
        finally:
            # 延迟清理临时文件，确保转换完成
            time.sleep(1)
            if os.path.exists(word_temp):
                try:
                    os.remove(word_temp)
                    print(f"Temporary Word file removed: {word_temp}")
                except Exception as e:
                    print(f"警告：无法删除临时文件 {word_temp}: {str(e)}")
    
    def _direct_pdf_generation(self, content, output_path):
        """直接生成PDF，作为备选方案"""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        
        try:
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
            
            # 获取样式表
            styles = getSampleStyleSheet()
            
            # 定义自定义样式
            title_style = ParagraphStyle(
                'TitleStyle',
                parent=styles['Heading1'],
                alignment=TA_CENTER,
                fontSize=16,
                spaceAfter=24
            )
            
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['BodyText'],
                alignment=TA_JUSTIFY,
                fontSize=12,
                spaceAfter=12,
                leading=18
            )
            
            # 处理内容，将换行符转换为Paragraph对象
            elements = []
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    elements.append(Spacer(1, 12))
                else:
                    elements.append(Paragraph(line, body_style))
            
            # 构建文档
            doc.build(elements)
            print(f"PDF saved successfully using direct method: {output_path}")
        except Exception as e:
            print(f"❌ 直接生成PDF也失败：{str(e)}")
            raise
    
    def generate_pdf(self, instruction, output_path="generated_document.pdf"):
        """完整流程：生成公文并保存为PDF"""
        # 生成文档
        content = self.generate_document(instruction)
        # 保存为PDF
        self.save_to_pdf(content, output_path)
        return output_path

import argparse

def main():
    # 添加命令行参数支持
    parser = argparse.ArgumentParser(description='卫健委公文生成器')
    parser.add_argument('--instruction', type=str, help='生成公文的指令')
    parser.add_argument('--output', type=str, default='generated_document.pdf', help='输出PDF文件路径')
    parser.add_argument('--interactive', action='store_true', help='以交互式模式运行')
    args = parser.parse_args()
    
    # 模型路径
    model_path = "./qwen2.5-7b-health-llamafactory"
    
    # 创建生成器实例
    generator = HealthDocumentGenerator(model_path)
    
    # 如果提供了指令，直接生成PDF
    if args.instruction:
        try:
            output_path = generator.generate_pdf(args.instruction, args.output)
            print(f"\n✅ 公文已成功生成！")
            print(f"📄 PDF文件路径：{output_path}\n")
        except Exception as e:
            print(f"❌ 生成失败：{str(e)}\n")
    else:
        # 交互式模式
        print("\n=== 卫健委公文生成器 ===")
        print("请输入您的指令，例如：'撰写一份关于开展医疗机构专项检查的通知'")
        print("输入 'exit' 退出程序\n")
        
        while True:
            # 接收用户输入
            instruction = input("请输入指令：")
            
            if instruction.lower() == 'exit':
                print("程序已退出")
                break
            
            if not instruction.strip():
                print("指令不能为空，请重新输入")
                continue
            
            try:
                # 生成PDF
                output_path = generator.generate_pdf(instruction)
                print(f"\n✅ 公文已成功生成！")
                print(f"📄 PDF文件路径：{output_path}\n")
            except Exception as e:
                print(f"❌ 生成失败：{str(e)}\n")

if __name__ == "__main__":
    main()